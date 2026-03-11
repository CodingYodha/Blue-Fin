import os
import json
import logging
import subprocess
import re as _re
from datetime import datetime
from typing import List, Dict, Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

from .context_assembler import assemble_cam_context, CAMContext
from .persona_accountant import run_accountant_persona
from .persona_compliance import run_compliance_persona
from .persona_cro import run_cro_persona

logger = logging.getLogger(__name__)


def _tex_escape(text: str) -> str:
    """Escape special LaTeX characters in text."""
    if not isinstance(text, str):
        text = str(text)
    replacements = {
        '\\': r'\textbackslash{}',
        '&': r'\&',
        '%': r'\%',
        '$': r'\$',
        '#': r'\#',
        '_': r'\_',
        '{': r'\{',
        '}': r'\}',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
        '<': r'\textless{}',
        '>': r'\textgreater{}',
        '₹': r'\rupee{}',
        '→': r'$\rightarrow$',
        '✅': r'\checkmark{}',
        '⚠': r'!',
        '✗': r'X',
    }
    # Process backslash first, then others
    text = text.replace('\\', '\x00BACKSLASH\x00')
    for char, replacement in replacements.items():
        if char == '\\':
            continue
        text = text.replace(char, replacement)
    text = text.replace('\x00BACKSLASH\x00', r'\textbackslash{}')
    return text


def _tex_wrap_long_text(text: str, max_line_len: int = 500) -> str:
    """Split very long paragraphs of persona output into LaTeX paragraphs."""
    escaped = _tex_escape(text)
    # Replace double newlines with LaTeX paragraph breaks
    escaped = escaped.replace('\n\n', '\n\n\\vspace{6pt}\n\n')
    # Replace single newlines with line breaks
    escaped = escaped.replace('\n', ' \\\\\n')
    return escaped

def build_audit_trail(
    accountant_output: dict,
    compliance_output: dict,
    cro_output: dict,
    ctx: CAMContext
) -> List[dict]:
    """
    Collects all source_citations from all three personas and adds
    standard system-generated entries for ML scores and stress tests.
    Returns deduplicated list ready for the appendix table.
    """
    citations = []

    # Collect from all persona outputs
    for output in [accountant_output, compliance_output, cro_output]:
        citations.extend(output.get("source_citations", []))

    # Add standard ML citations
    citations.append({
        "claim": f"Financial Health score {ctx.score_financial:.1f}/40",
        "source": "LightGBM Model 1 — trained on 5000 synthetic companies calibrated to CRISIL benchmarks",
        "module": "Core ML Engine — Section 7",
        "confidence": "HIGH"
    })
    citations.append({
        "claim": f"Sector sentiment {ctx.sector_sentiment_score:.2f} ({ctx.sector_risk})",
        "source": "Claude sentiment scoring with Indian regulatory severity mappings (V3 fix)",
        "module": "LangGraph Research Agent — Section 6",
        "confidence": "HIGH"
    })

    # Add stress test citations
    for scenario, result in ctx.stress_summary.items():
        if isinstance(result, dict):
            citations.append({
                "claim": f"Stress test {scenario}: score → {result.get('stressed_score', 0):.1f}",
                "source": "LightGBM stress scenario engine — Section 7.8",
                "module": "Core ML Engine",
                "confidence": "HIGH"
            })

    # Deduplicate by claim text
    seen = set()
    deduped = []
    for c in citations:
        claim_text = c.get("claim", "")
        if claim_text not in seen:
            seen.add(claim_text)
            deduped.append(c)

    return deduped


def generate_latex_pdf(cam_data: dict, job_id: str):
    """Generates a professional Credit Appraisal Memo as LaTeX → PDF."""
    base_dir = f"/tmp/intelli-credit/{job_id}"
    os.makedirs(base_dir, exist_ok=True)

    company = _tex_escape(cam_data.get('company_name', 'Company'))
    cin = _tex_escape(cam_data.get('company_cin', ''))
    sector = _tex_escape(cam_data.get('sector', ''))
    promoters = _tex_escape(', '.join(cam_data.get('promoters', [])))
    loan_amt = cam_data.get('loan_amount', 0.0)
    date_str = cam_data.get('date', datetime.utcnow().strftime('%Y-%m-%d'))

    cro_out = cam_data.get("cro_output", {})
    acc_out = cam_data.get("accountant_output", {})
    comp_out = cam_data.get("compliance_output", {})
    final_dec = cro_out.get("final_decision", "UNKNOWN")

    dec_color = "red" if final_dec == "REJECT" else ("green!60!black" if final_dec == "APPROVE" else "orange")

    # Build score table rows
    score_rows = [
        ("Financial Health (M1)", f"{cam_data.get('score_financial', 0)}", "40"),
        ("Credit Behaviour (M2)", f"{cam_data.get('score_behaviour', 0)}", "30"),
        ("External Risk (M3)", f"{cam_data.get('score_external', 0)}", "20"),
        ("Text Signals (M4)", f"{cam_data.get('score_text', 0)}", "10"),
        ("Layer 1 — Rules", f"{cam_data.get('layer1_score', 0)}", "100"),
        ("Layer 2 — ML", f"{cam_data.get('layer2_score', 0)}", "100"),
    ]
    score_table_rows = "\n".join(
        f"        {_tex_escape(n)} & {s} & {m} \\\\ \\hline"
        for n, s, m in score_rows
    )

    # Build SHAP rows
    shap_rows_tex = ""
    for model_name, drivers in cam_data.get('shap_by_model', {}).items():
        for sh in drivers[:3]:
            direction = sh.get('direction', '')
            label = _tex_escape(sh.get('human_label', sh.get('feature', '')))
            val = sh.get('shap_value', 0)
            dir_symbol = r"$\uparrow$" if direction == "risk_increasing" else r"$\downarrow$"
            shap_rows_tex += f"        {_tex_escape(model_name)} & {label} & {val:.4f} & {dir_symbol} \\\\ \\hline\n"

    # Build audit trail rows
    audit_rows_tex = ""
    for c in cam_data.get("audit_trail", []):
        claim = _tex_escape(c.get('claim', ''))[:80]
        source = _tex_escape(c.get('source', ''))[:80]
        module = _tex_escape(c.get('module', ''))
        conf = c.get('confidence', 'UNKNOWN')
        audit_rows_tex += f"        {claim} & {source} & {module} & {conf} \\\\ \\hline\n"

    # Build covenants
    covenants_tex = ""
    mand_cov = cro_out.get("mandatory_covenants", [])
    if mand_cov:
        covenants_tex += "\\subsection*{Mandatory Covenants}\n\\begin{itemize}\n"
        for mc in mand_cov:
            covenants_tex += f"  \\item {_tex_escape(str(mc))}\n"
        covenants_tex += "\\end{itemize}\n"

    mon_trig = cro_out.get("monitoring_triggers", [])
    if mon_trig:
        covenants_tex += "\\subsection*{Monitoring Triggers}\n\\begin{itemize}\n"
        for mt in mon_trig:
            covenants_tex += f"  \\item {_tex_escape(str(mt))}\n"
        covenants_tex += "\\end{itemize}\n"

    # Officer notes section
    officer_notes_tex = ""
    off_notes = cam_data.get("officer_notes_text")
    if off_notes:
        officer_notes_tex = f"""
\\section{{Officer Field Visit Notes}}
{_tex_wrap_long_text(off_notes)}
"""
        if cam_data.get("injection_detected"):
            officer_notes_tex += "\n\\textbf{\\textcolor{red}{PROMPT INJECTION ATTEMPT DETECTED --- see audit log}}\n"

    # Sanctioned limit section
    sanction_tex = ""
    if final_dec in ["APPROVE", "CONDITIONAL"]:
        limit = cro_out.get('sanctioned_limit_crore', 0) or 0
        rate = cro_out.get('interest_rate_pct', 0) or 0
        sanction_tex = f"""
Sanctioned Limit: \\textbf{{Rs. {limit:.2f} Cr}} \\\\
Interest Rate: \\textbf{{{rate}\\%}}
"""

    override_tex = ""
    if cro_out.get("override_applied"):
        override_tex = f"""
\\begin{{tcolorbox}}[colback=red!5!white,colframe=red!75!black,title=OVERRIDE APPLIED]
{_tex_escape(cro_out.get('override_reason', ''))}
\\end{{tcolorbox}}
"""

    latex_src = r"""\documentclass[11pt,a4paper]{article}
\usepackage[margin=2.5cm]{geometry}
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
\usepackage{booktabs,tabularx,longtable}
\usepackage[table]{xcolor}
\usepackage{titlesec}
\usepackage{fancyhdr}
\usepackage{graphicx}
\usepackage{amssymb}
\usepackage{tcolorbox}
\usepackage{enumitem}
\usepackage{parskip}
\usepackage{hyperref}

\definecolor{icblue}{RGB}{37,99,235}
\definecolor{icgray}{RGB}{100,100,110}

\newcommand{\rupee}{Rs.}

\hypersetup{colorlinks=true,linkcolor=icblue,urlcolor=icblue}

\titleformat{\section}{\Large\bfseries\color{icblue}}{}{0em}{}[\titlerule]
\titleformat{\subsection}{\large\bfseries\color{icblue!70!black}}{}{0em}{}

\pagestyle{fancy}
\fancyhf{}
\fancyhead[L]{\small\textcolor{icgray}{IntelliCredit --- Credit Appraisal Memo}}
\fancyhead[R]{\small\textcolor{icgray}{""" + _tex_escape(date_str) + r"""}}
\fancyfoot[C]{\thepage}
\renewcommand{\headrulewidth}{0.4pt}

\begin{document}

% ── Title Page ──────────────────────────────────────────────
\begin{titlepage}
\centering
\vspace*{3cm}
{\Huge\bfseries\textcolor{icblue}{Credit Appraisal Memo}\par}
\vspace{1cm}
{\LARGE\bfseries """ + company + r"""\par}
\vspace{0.5cm}
{\large CIN: """ + cin + r""" \quad $\vert$ \quad Sector: """ + sector + r"""\par}
\vspace{0.5cm}
{\large Promoters: """ + promoters + r"""\par}
\vspace{0.5cm}
{\large Loan Requested: \textbf{\rupee{} """ + f"{loan_amt:.2f}" + r""" Cr}\par}
\vspace{1cm}
{\large Assessment Date: """ + _tex_escape(date_str) + r"""\par}
\vspace{2cm}
\begin{tcolorbox}[colback=""" + dec_color + r"""!8!white,colframe=""" + dec_color + r""",width=8cm,center]
\centering
{\LARGE\bfseries\textcolor{""" + dec_color + r"""}{""" + _tex_escape(final_dec) + r"""}}\\[6pt]
{\large Final Score: \textbf{""" + f"{cam_data.get('final_score', 0):.1f}" + r"""/100}}\\[4pt]
{\normalsize PD (Meta): """ + f"{cam_data.get('pd_meta', 0)*100:.2f}" + r"""\%}
\end{tcolorbox}
\vfill
{\small\textcolor{icgray}{Generated by IntelliCredit AI Engine --- Confidential}}
\end{titlepage}

\tableofcontents
\newpage

% ── Executive Summary ──────────────────────────────────────
\section{Executive Summary}

\begin{tcolorbox}[colback=""" + dec_color + r"""!5!white,colframe=""" + dec_color + r"""]
\textbf{Decision:} \textcolor{""" + dec_color + r"""}{""" + _tex_escape(final_dec) + r"""}
\quad\quad Final Score: \textbf{""" + f"{cam_data.get('final_score', 0):.1f}" + r"""/100}
\quad\quad PD: """ + f"{cam_data.get('pd_meta', 0)*100:.2f}" + r"""\%
""" + sanction_tex + r"""
\end{tcolorbox}
""" + override_tex + r"""

% ── Financial Assessment ───────────────────────────────────
\section{Financial Assessment}
""" + _tex_wrap_long_text(acc_out.get("content", "Financial assessment not available.")) + r"""

% ── Legal and Governance ───────────────────────────────────
\section{Legal and Governance Assessment}
""" + _tex_wrap_long_text(comp_out.get("content", "Compliance assessment not available.")) + r"""

% ── CRO Recommendation ────────────────────────────────────
\section{Chief Risk Officer Recommendation}
""" + _tex_wrap_long_text(cro_out.get("content", "CRO recommendation not available.")) + r"""

""" + covenants_tex + r"""

""" + officer_notes_tex + r"""

% ── ML Score Dashboard ────────────────────────────────────
\section{ML Score Dashboard}

\begin{center}
\begin{tabularx}{0.75\textwidth}{|X|c|c|}
\hline
\rowcolor{icblue!15}
\textbf{Component} & \textbf{Score} & \textbf{Max} \\ \hline
""" + score_table_rows + r"""
\end{tabularx}
\end{center}

\subsection*{SHAP Top Drivers}
\begin{center}
\begin{tabularx}{\textwidth}{|l|X|c|c|}
\hline
\rowcolor{icblue!15}
\textbf{Model} & \textbf{Feature} & \textbf{SHAP Value} & \textbf{Direction} \\ \hline
""" + shap_rows_tex + r"""
\end{tabularx}
\end{center}

% ── Audit Trail ───────────────────────────────────────────
\section{Audit Trail --- Source Citation Table}

\begin{center}
\small
\begin{longtable}{|p{4cm}|p{4.5cm}|p{3cm}|c|}
\hline
\rowcolor{icblue!15}
\textbf{CAM Claim} & \textbf{Source} & \textbf{Module} & \textbf{Confidence} \\ \hline
\endhead
""" + audit_rows_tex + r"""
\end{longtable}
\end{center}

\vfill
\begin{center}
\small\textcolor{icgray}{--- End of Credit Appraisal Memo ---}
\end{center}

\end{document}
"""

    tex_path = os.path.join(base_dir, "cam_final.tex")
    with open(tex_path, "w", encoding="utf-8") as f:
        f.write(latex_src)
    logger.info(f"Saved LaTeX source to {tex_path}")

    # Compile to PDF with pdflatex (run twice for TOC)
    try:
        for _ in range(2):
            subprocess.run(
                ["pdflatex", "-interaction=nonstopmode", "-output-directory", base_dir, tex_path],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,
            )
        logger.info(f"Generated LaTeX PDF for job {job_id}")
    except FileNotFoundError:
        logger.warning("pdflatex not found. Install texlive-latex-base. Skipping PDF generation.")
    except subprocess.CalledProcessError as e:
        err_msg = ""
        if e.stdout:
            err_msg = e.stdout.decode(errors='replace')[-800:]
        elif e.stderr:
            err_msg = e.stderr.decode(errors='replace')[-500:]
        else:
            err_msg = "unknown error"
        logger.warning(f"pdflatex compilation failed: {err_msg}")
    except subprocess.TimeoutExpired:
        logger.warning("pdflatex timed out after 60s")


def generate_docx(cam_data: dict, job_id: str):
    """Generates the Word document using python-docx (fallback)."""
    doc = Document()
    
    # Header
    head = doc.add_heading(cam_data.get('company_name', 'Company Name'), level=0)
    for run in head.runs:
        run.bold = True
        
    doc.add_paragraph(f"CIN: {cam_data.get('company_cin', '')} | Sector: {cam_data.get('sector', '')}")
    doc.add_paragraph(f"Promoters: {', '.join(cam_data.get('promoters', []))}")
    doc.add_paragraph(f"Loan Requested: ₹{cam_data.get('loan_amount', 0.0):.2f} Cr")
    doc.add_paragraph(f"Assessment Date: {cam_data.get('date', datetime.utcnow().strftime('%Y-%m-%d'))}")
    
    doc.add_page_break()

    # SECTION 2: EXECUTIVE SUMMARY
    doc.add_heading("2. Executive Summary", level=1)
    
    cro_out = cam_data.get("cro_output", {})
    final_dec = cro_out.get("final_decision", "UNKNOWN")
    
    p = doc.add_paragraph("Final Decision: ")
    r = p.add_run(final_dec)
    r.bold = True
    if final_dec == "APPROVE":
        r.font.color.rgb = RGBColor(0, 128, 0)
    elif final_dec == "REJECT":
        r.font.color.rgb = RGBColor(255, 0, 0)
    else:
        r.font.color.rgb = RGBColor(255, 165, 0)
        
    doc.add_paragraph(f"Final Score: {cam_data.get('final_score', 0):.1f}/100")
    doc.add_paragraph(f"Probability of Default (Meta): {cam_data.get('pd_meta', 0)*100:.2f}%")
    
    if final_dec in ["APPROVE", "CONDITIONAL"]:
        doc.add_paragraph(f"Sanctioned Limit: ₹{cro_out.get('sanctioned_limit_crore', 0) or 0:.2f} Cr")
        doc.add_paragraph(f"Interest Rate: {cro_out.get('interest_rate_pct', 0) or 0}%")

    if cro_out.get("override_applied"):
        p_over = doc.add_paragraph("OVERRIDE APPLIED")
        for run in p_over.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)
        doc.add_paragraph(f"Reason: {cro_out.get('override_reason', '')}")

    doc.add_page_break()

    # SECTION 3: FINANCIAL ASSESSMENT
    doc.add_heading("3. Financial Assessment", level=1)
    acc_out = cam_data.get("accountant_output", {})
    doc.add_paragraph(acc_out.get("content", "Financial assessment not available."))

    # Note: Confidence badges formatting could be post-processed here if text contains [✅ HIGH CONFIDENCE]
    # In a full implementation, we'd regex search and colorize the runs.
    
    # SECTION 4: LEGAL AND GOVERNANCE
    doc.add_heading("4. Legal and Governance Assessment", level=1)
    comp_out = cam_data.get("compliance_output", {})
    doc.add_paragraph(comp_out.get("content", "Compliance assessment not available."))
    
    # SECTION 5: CHIEF RISK OFFICER RECOMMENDATION
    doc.add_heading("5. Chief Risk Officer Recommendation", level=1)
    doc.add_paragraph(cro_out.get("content", "CRO recommendation not available."))
    
    mand_cov = cro_out.get("mandatory_covenants", [])
    if mand_cov:
        doc.add_heading("Mandatory Covenants", level=2)
        for mc in mand_cov:
            doc.add_paragraph(str(mc), style='List Bullet')
            
    mon_trig = cro_out.get("monitoring_triggers", [])
    if mon_trig:
        doc.add_heading("Monitoring Triggers", level=2)
        for mt in mon_trig:
            doc.add_paragraph(str(mt), style='List Bullet')

    doc.add_page_break()

    # SECTION 6: OFFICER FIELD VISIT NOTES
    off_notes = cam_data.get("officer_notes_text")
    if off_notes:
        doc.add_heading("6. Officer Field Visit Notes", level=1)
        doc.add_paragraph(f"Notes:\n{off_notes}")
        
        if cam_data.get("injection_detected"):
            p_inj = doc.add_paragraph("PROMPT INJECTION ATTEMPT DETECTED — see audit log")
            for run in p_inj.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                
        doc.add_paragraph(f"Score Adjustments Applied: {cam_data.get('officer_notes_adj', {})}")

    doc.add_page_break()

    # SECTION 7: ML SCORE DASHBOARD (Table)
    doc.add_heading("7. ML Score Dashboard", level=1)
    table_scores = doc.add_table(rows=1, cols=2)
    table_scores.style = 'Table Grid'
    hdr_cells = table_scores.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Score'
    
    components = [
        ("Financial Health", f"{cam_data.get('score_financial', 0)}/40"),
        ("Credit Behaviour", f"{cam_data.get('score_behaviour', 0)}/30"),
        ("External Risk", f"{cam_data.get('score_external', 0)}/20"),
        ("Text Signals", f"{cam_data.get('score_text', 0)}/10"),
        ("Layer 1 (Rules)", f"{cam_data.get('layer1_score', 0)}/100"),
        ("Layer 2 (ML)", f"{cam_data.get('layer2_score', 0)}/100")
    ]
    for name, val in components:
        row_cells = table_scores.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = val

    doc.add_heading("SHAP Top Drivers", level=2)
    for k, v_list in cam_data.get('shap_by_model', {}).items():
        doc.add_paragraph(str(k), style='Heading 3')
        for sh in v_list[:3]:
            doc.add_paragraph(f"{sh.get('direction')}: {sh.get('human_label')} ({sh.get('shap_value')})", style='List Bullet')

    doc.add_page_break()

    # SECTION 8: AUDIT TRAIL APPENDIX
    doc.add_heading("8. Audit Trail Appendix — Source Citation Table", level=1)
    table_audit = doc.add_table(rows=1, cols=4)
    table_audit.style = 'Table Grid'
    hdr_cells = table_audit.rows[0].cells
    hdr_cells[0].text = 'CAM Claim'
    hdr_cells[1].text = 'Source'
    hdr_cells[2].text = 'Module'
    hdr_cells[3].text = 'Confidence'

    for citation in cam_data.get("audit_trail", []):
        row_cells = table_audit.add_row().cells
        row_cells[0].text = citation.get('claim', '')
        row_cells[1].text = citation.get('source', '')
        row_cells[2].text = citation.get('module', '')
        row_cells[3].text = citation.get('confidence', 'UNKNOWN')

    base_dir = f"/tmp/intelli-credit/{job_id}"
    os.makedirs(base_dir, exist_ok=True)
    docx_path = os.path.join(base_dir, "cam_final.docx")
    doc.save(docx_path)
    logger.info(f"Saved DOCX to {docx_path}")


async def generate_cam_pipeline(job_id: str):
    """
    Main background pipeline:
    1. Assemble Context
    2. Run P1 (Accountant)
    3. Run P2 (Compliance)
    4. Run P3 (CRO)
    5. Build Audit Trail
    6. Save Draft JSON
    7. Generate DOCX/PDF
    """
    logger.info(f"Starting CAM Generation for job {job_id}")
    
    ctx = await assemble_cam_context(job_id)
    
    import asyncio
    
    # 1 & 2. Concurrent Accountant and Compliance Officer Runs
    acc_out, comp_out = await asyncio.gather(
        run_accountant_persona(ctx),
        run_compliance_persona(ctx)
    )
    
    # 3. Chief Risk Officer
    cro_out = await run_cro_persona(ctx, acc_out, comp_out)
    
    # 4. Audit Trail
    audit_trail = build_audit_trail(acc_out, comp_out, cro_out, ctx)
    
    # 5. Build combined data dictionary
    cam_data = {
        "job_id": job_id,
        "company_name": ctx.company_name,
        "company_cin": ctx.company_cin,
        "sector": ctx.sector,
        "promoters": ctx.promoter_names,
        "loan_amount": ctx.loan_amount_requested,
        "date": datetime.utcnow().isoformat(),
        
        "final_score": ctx.final_score,
        "layer1_score": ctx.layer1_score,
        "layer2_score": ctx.layer2_score,
        "pd_meta": ctx.pd_meta,
        "score_financial": ctx.score_financial,
        "score_behaviour": ctx.score_behaviour,
        "score_external": ctx.score_external,
        "score_text": ctx.score_text,
        
        "accountant_output": acc_out,
        "compliance_output": comp_out,
        "cro_output": cro_out,
        
        "officer_notes_text": ctx.officer_notes_text,
        "officer_notes_adj": ctx.officer_notes_adj,
        "injection_detected": ctx.injection_detected,
        
        "shap_by_model": ctx.shap_by_model,
        "audit_trail": audit_trail
    }
    
    base_dir = f"/tmp/intelli-credit/{job_id}"
    os.makedirs(base_dir, exist_ok=True)
    
    draft_path = os.path.join(base_dir, "cam_draft.json")
    with open(draft_path, "w") as f:
        json.dump(cam_data, f, indent=2)
        
    logger.info(f"Saved CAM Draft JSON to {draft_path}")
    
    # 6. Generate LaTeX PDF (primary) and DOCX (fallback)
    generate_latex_pdf(cam_data, job_id)
    generate_docx(cam_data, job_id)
    
    logger.info(f"CAM Generation COMPLETE for job {job_id}")
